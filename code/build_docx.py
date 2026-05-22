# -*- coding: utf-8 -*-
"""build_docx.py — paper_b5.html 내용을 학과 양식(B5)에 맞춘 정밀 DOCX로 변환.
- 용지 B5(182x257), 여백 상하15/좌우25mm
- 본문 바탕 10.5pt, 줄간격 160%
- front matter(표지/속표지/인준서/차례/국문요약) + 십진 본문 + 표6 + 그림5 + 참고문헌
출력: 임창우_졸업논문(docx버전).docx
"""
import os, re
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
KOR = "바탕"

soup = BeautifulSoup(open(os.path.join(BASE, "paper_b5.html"), encoding="utf-8").read(), "html.parser")

doc = Document()

# ---------- 페이지/기본 스타일 ----------
sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(182), Mm(257)
sec.top_margin = sec.bottom_margin = Mm(15)
sec.left_margin = sec.right_margin = Mm(25)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(10.5)
normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), KOR)
pf = normal.paragraph_format
pf.line_spacing = 1.6
pf.space_after = Pt(0)


def krun(p, text, size=10.5, bold=False, italic=False, color=None, sub=False, sup=False):
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), KOR)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    if sub: r.font.subscript = True
    if sup: r.font.superscript = True
    return r


def render_inline(p, node, bold=False, italic=False, sub=False, sup=False, size=10.5):
    for ch in node.children:
        if isinstance(ch, NavigableString):
            txt = re.sub(r"\s+", " ", str(ch))
            if txt:
                krun(p, txt, size=size, bold=bold, italic=italic, sub=sub, sup=sup)
        else:
            nm = (ch.name or "").lower()
            if nm == "br":
                p.add_run().add_break(); continue
            render_inline(p, ch,
                          bold or nm in ("b", "strong"),
                          italic or nm in ("i", "em"),
                          sub or nm == "sub",
                          sup or nm == "sup", size)


def heading(text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(2)
    krun(p, text, size=14 if level == 1 else 11.5, bold=True)
    return p


def center_blank(n=1):
    for _ in range(n):
        doc.add_paragraph()


def add_page_field(p):
    """문단에 현재 페이지 번호 필드(PAGE) 삽입."""
    r = p.add_run(); r.font.size = Pt(10); r.font.name = "Times New Roman"
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), KOR)
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    r._r.append(f1); r._r.append(it); r._r.append(f2)


# ================= FRONT MATTER =================
TKO1, TKO2 = "4종 멤리스터 소재 파라미터를 활용한", "PIM 환경 웨이퍼 결함 분류 성능 비교 연구"
TEN1 = "A Comparative Study of Four Memristive Material Candidates"
TEN2 = "for Processing-In-Memory Wafer Defect Classification"

def C(text, size=10.5, bold=False, italic=False, after=0, before=0):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    krun(p, text, size=size, bold=bold, italic=italic); return p

# 표지
center_blank(2)
C("학사학위논문", 15, bold=True, after=6)
center_blank(3)
C(TKO1, 17, bold=True, after=2); C(TKO2, 17, bold=True, after=10)
C(TEN1, 11, italic=True, after=1); C(TEN2, 11, italic=True)
center_blank(5)
C("상명대학교 공과대학", 12, after=4)
C("시스템반도체공학과", 12, after=4)
C("임 창 우", 12, after=12)
C("2026년 5월", 12)
doc.add_page_break()

# 속표지
center_blank(2)
C(TKO1, 16, bold=True, after=2); C(TKO2, 16, bold=True, after=8)
C(TEN1, 11, italic=True, after=1); C(TEN2, 11, italic=True, after=14)
C("지도교수　이 종 환", 12, after=14)
C("본 논문을 학사학위 논문으로 제출함", 11, after=16)
C("상명대학교 공과대학", 12, after=4)
C("시스템반도체공학과", 12, after=4)
C("임 창 우", 12, after=12)
C("2026년 5월", 12)
doc.add_page_break()

# 인준서
center_blank(3)
C("임 창 우의 학사학위 논문을", 15, bold=True, after=2)
C("인준함", 15, bold=True, after=20)
for role in ["심사위원장", "심사위원", "심사위원"]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    krun(p, f"{role}　　　　　　　　　　　(인)", size=12)
center_blank(2)
C("상명대학교 공과대학", 12, after=6)
C("2026년 5월", 12)
doc.add_page_break()

# 차례
C("차　　례", 14, bold=True, after=12)
TOC = [("국문 요약", 0), ("1. 서론", 0), ("1.1 연구 배경", 1),
       ("1.2 본 연구의 기여", 1), ("2. 이론적 배경", 0), ("2.1 Processing-In-Memory (PIM)", 1),
       ("2.2 멤리스터 크로스바 어레이", 1), ("2.3 VTEAM 모델", 1), ("2.4 웨이퍼 결함 분류와 CNN", 1),
       ("3. 방법론", 0), ("3.1 데이터셋", 1), ("3.2 CNN 모델 구조", 1), ("3.3 4종 소재 파라미터", 1),
       ("3.4 매핑·양자화 및 비이상성 모델", 1), ("3.5 평가 프로토콜", 1), ("4. 결과", 0),
       ("4.1 Baseline (FP32 software) 성능", 1), ("4.2 4종 소재 Ideal vs Realistic 비교", 1),
       ("4.3 Worst-case 붕괴 현상 분석", 1), ("4.4 결함 유형별 (per-class) F1 분석", 1),
       ("5. 논의", 0), ("5.1 저항비의 영향", 1), ("5.2 Metallic Nanowire의 본질적 한계", 1),
       ("5.3 Worst-case 붕괴 현상의 PIM 함의", 1), ("5.4 한계 및 향후 과제", 1),
       ("6. 결론", 0), ("참고문헌", 0)]
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
PAGES = {  # 차례 페이지 번호 (국문요약=1 기준; PDF 렌더로 매핑)
    "국문 요약": 1, "1. 서론": 2, "1.1 연구 배경": 2, "1.2 본 연구의 기여": 2,
    "2. 이론적 배경": 3, "2.1 Processing-In-Memory (PIM)": 3, "2.2 멤리스터 크로스바 어레이": 3,
    "2.3 VTEAM 모델": 3, "2.4 웨이퍼 결함 분류와 CNN": 4, "3. 방법론": 4, "3.1 데이터셋": 4,
    "3.2 CNN 모델 구조": 5, "3.3 4종 소재 파라미터": 6, "3.4 매핑·양자화 및 비이상성 모델": 6,
    "3.5 평가 프로토콜": 8, "4. 결과": 8, "4.1 Baseline (FP32 software) 성능": 8,
    "4.2 4종 소재 Ideal vs Realistic 비교": 8, "4.3 Worst-case 붕괴 현상 분석": 9,
    "4.4 결함 유형별 (per-class) F1 분석": 10, "5. 논의": 11, "5.1 저항비의 영향": 11,
    "5.2 Metallic Nanowire의 본질적 한계": 12, "5.3 Worst-case 붕괴 현상의 PIM 함의": 12,
    "5.4 한계 및 향후 과제": 12, "6. 결론": 13, "참고문헌": 14,
}
for txt, lv in TOC:
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Mm(6 * lv)
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.tab_stops.add_tab_stop(Mm(132 - 6 * lv), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    disp = re.sub(r"^(\d+\.\d+)(\s)", r"\1.\2", txt)
    krun(p, disp, size=10.5, bold=(lv == 0))
    if txt in PAGES:
        krun(p, "\t" + str(PAGES[txt]), size=10.5, bold=(lv == 0))

# === 본문 섹션 시작(국문요약부터 페이지 번호 1) ===
bsec = doc.add_section(WD_SECTION.NEW_PAGE)
bsec.page_width, bsec.page_height = Mm(182), Mm(257)
bsec.top_margin = bsec.bottom_margin = Mm(15); bsec.left_margin = bsec.right_margin = Mm(25)
ftr = bsec.footer; ftr.is_linked_to_previous = False
fp = ftr.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_page_field(fp)
_pg = OxmlElement("w:pgNumType"); _pg.set(qn("w:start"), "1"); bsec._sectPr.append(_pg)

# 국문 요약
C("국 문 요 약", 14, bold=True, after=10)
summary = soup.find("section", class_="summary")
for p_el in summary.find_all("p"):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    render_inline(p, p_el)
doc.add_page_break()


# ================= 본문 =================
FIG_W = {  # 그림별 폭(mm) — B5 분량 절감 위해 축소
    "wafer_defect_examples.png": 70, "cnn_architecture.png": 100,
    "pipeline_overview.png": 100, "material_comparison.png": 82,
    "bimodal_distribution.png": 82,
}

def add_table(tbl_el):
    rows = tbl_el.find_all("tr")
    has_span = bool(tbl_el.find(attrs={"rowspan": True})) or bool(tbl_el.find(attrs={"colspan": True}))
    if has_span:  # Table 4 — 평탄화된 9열 헤더로 처리
        header = ["Material", "Roff/Ron", "States", "Ideal Subset Acc", "Ideal Macro F1",
                  "Realistic Subset Acc", "Realistic Macro F1", "ΔF1 abs", "ΔF1 rel"]
        data = [r for r in rows if r.find("td")]
        t = doc.add_table(rows=1, cols=len(header)); t.style = "Table Grid"
        for j, h in enumerate(header):
            c = t.rows[0].cells[j]; c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            krun(c.paragraphs[0], h, size=7.5, bold=True)
        for r in data:
            cells = r.find_all("td"); row = t.add_row().cells
            for j, cell in enumerate(cells):
                pp = row[j].paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                render_inline(pp, cell, size=8);
                for run in pp.runs: run.font.size = Pt(8)
    else:
        ncol = max(len(r.find_all(["th", "td"])) for r in rows)
        t = doc.add_table(rows=0, cols=ncol); t.style = "Table Grid"
        for r in rows:
            cells = r.find_all(["th", "td"]); is_h = bool(r.find("th"))
            row = t.add_row().cells
            for j, cell in enumerate(cells):
                if j >= ncol: break
                pp = row[j].paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                render_inline(pp, cell, bold=is_h, size=9)
                for run in pp.runs: run.font.size = Pt(9); run.bold = is_h
    t.alignment = WD_TABLE_ALIGNMENT.CENTER


main = soup.find("main", class_="flow")
for el in main.find_all(recursive=False):
    nm = el.name
    cls = el.get("class", [])
    if nm == "h1":
        heading(el.get_text(strip=True), 1)
    elif nm == "h2":
        heading(re.sub(r"^(\d+\.\d+)(\s)", r"\1.\2", el.get_text(strip=True)), 2)
    elif nm == "table":
        add_table(el)
    elif nm == "div" and "figure" in cls:
        img = el.find("img")
        if img:
            fn = os.path.basename(img["src"])
            path = os.path.join(BASE, "figures", fn)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run()
            run.add_picture(path, width=Mm(FIG_W.get(fn, 120)))
        cap = el.find("div", class_="caption")
        if cap:
            cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(5); cp.paragraph_format.line_spacing = 1.2
            render_inline(cp, cap, italic=True, size=9)
    elif nm == "div" and "equation" in cls:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        render_inline(p, el, italic=True)
    elif nm == "div" and "references" in cls:
        ol = el.find("ol")
        for i, li in enumerate(ol.find_all("li"), 1):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Mm(5)
            p.paragraph_format.first_line_indent = Mm(-5); p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.25
            krun(p, f"[{i}] ", size=8.5)
            render_inline(p, li, size=8.5)
            for r in p.runs: r.font.size = Pt(8.5)
    elif nm == "p":
        p = doc.add_paragraph()
        if "caption" in cls:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; render_inline(p, el, italic=True, size=9)
            p.paragraph_format.space_after = Pt(8)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if "noindent" not in cls:
                p.paragraph_format.first_line_indent = Pt(11)
            render_inline(p, el)
            p.paragraph_format.space_after = Pt(2)
    elif nm == "ul":
        for li in el.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Bullet"); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(2)
            render_inline(p, li)
    elif nm == "h1" or nm == "h3":
        heading(el.get_text(strip=True), 1)

out = os.path.join(BASE, "임창우_졸업논문(docx).docx")
doc.save(out)
print("저장:", out, os.path.getsize(out), "bytes")
